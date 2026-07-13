from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_db
from app.models import Chapter, Project, Volume

router = APIRouter()


class ExportRequest(BaseModel):
    project_id: str


async def _get_export_data(db: AsyncSession, project_id: str):
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    volumes = (await db.execute(
        select(Volume).where(Volume.project_id == project_id).order_by(Volume.sort_order)
    )).scalars().all()

    chapters = (await db.execute(
        select(Chapter)
        .join(Volume, Chapter.volume_id == Volume.id)
        .where(Chapter.project_id == project_id)
        .order_by(Volume.sort_order, Chapter.sort_order)
    )).scalars().all()

    return project, volumes, chapters


@router.post("/txt")
async def export_txt(req: ExportRequest, db: AsyncSession = Depends(get_db)):
    project, volumes, chapters = await _get_export_data(db, req.project_id)

    lines = [project.name, "=" * len(project.name) * 2, ""]

    vol_map: dict[str, list] = {}
    for ch in chapters:
        vol_map.setdefault(ch.volume_id, []).append(ch)

    for vol in volumes:
        lines.append(f"\n{'=' * 40}")
        lines.append(vol.title)
        lines.append(f"{'=' * 40}\n")
        for ch in vol_map.get(vol.id, []):
            lines.append(f"\n{ch.title}")
            lines.append("-" * 30)
            lines.append(ch.content or "(暂无内容)")
            lines.append("")

    content = "\n".join(lines)
    return Response(
        content=content.encode("utf-8"),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(project.name)}.txt"},
    )


@router.post("/markdown")
async def export_markdown(req: ExportRequest, db: AsyncSession = Depends(get_db)):
    project, volumes, chapters = await _get_export_data(db, req.project_id)

    lines = [f"# {project.name}", ""]

    vol_map: dict[str, list] = {}
    for ch in chapters:
        vol_map.setdefault(ch.volume_id, []).append(ch)

    for vol in volumes:
        lines.append(f"\n## {vol.title}\n")
        for ch in vol_map.get(vol.id, []):
            lines.append(f"### {ch.title}\n")
            lines.append(ch.content or "*暂无内容*")
            lines.append("")

    content = "\n".join(lines)
    return Response(
        content=content.encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(project.name)}.md"},
    )


@router.post("/docx")
async def export_docx(req: ExportRequest, db: AsyncSession = Depends(get_db)):
    from docx import Document
    from io import BytesIO

    project, volumes, chapters = await _get_export_data(db, req.project_id)

    doc = Document()
    doc.add_heading(project.name, 0)

    vol_map: dict[str, list] = {}
    for ch in chapters:
        vol_map.setdefault(ch.volume_id, []).append(ch)

    for vol in volumes:
        doc.add_heading(vol.title, 1)
        for ch in vol_map.get(vol.id, []):
            doc.add_heading(ch.title, 2)
            for para in (ch.content or "").split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        content=buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(project.name)}.docx"},
    )


@router.post("/epub")
async def export_epub(req: ExportRequest, db: AsyncSession = Depends(get_db)):
    from ebooklib import epub
    from io import BytesIO

    project, volumes, chapters = await _get_export_data(db, req.project_id)

    book = epub.EpubBook()
    book.set_identifier(project.id)
    book.set_title(project.name)
    book.set_language("zh")

    vol_map: dict[str, list] = {}
    for ch in chapters:
        vol_map.setdefault(ch.volume_id, []).append(ch)

    spine = ["nav"]
    toc = []

    for vi, vol in enumerate(volumes):
        vol_chapters = []
        for ci, ch in enumerate(vol_map.get(vol.id, [])):
            content_html = "".join(f"<p>{p.strip()}</p>" for p in (ch.content or "").split("\n") if p.strip())
            epub_ch = epub.EpubHtml(title=ch.title, file_name=f"ch_{vi}_{ci}.xhtml", lang="zh")
            epub_ch.content = f"<h2>{ch.title}</h2>{content_html}"
            book.add_item(epub_ch)
            spine.append(epub_ch)
            vol_chapters.append(epub_ch)
        if vol_chapters:
            toc.append((epub.Section(vol.title), vol_chapters))

    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    buffer = BytesIO()
    epub.write_epub(buffer, book)
    buffer.seek(0)

    return Response(
        content=buffer.read(),
        media_type="application/epub+zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(project.name)}.epub"},
    )
